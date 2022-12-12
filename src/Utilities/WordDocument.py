from typing import List

import docx
from docx import Document
import os

from docx.shared import Inches
from docx.table import Table

from Models.Components.AnswerChoice import AnswerChoice
from Models.Components.Emphasis import Emphasis
from Models.Components.Graphic import Graphic
from Models.Components.MultiYesNoAnswer import MultiYesNoAnswer
from Models.Questions.FiveAnswerQuestion import FiveAnswerQuestion
from Models.Questions.SingleAnswerQuestion import SingleAnswerQuestion


class WordDocument:

    @staticmethod
    def get_document_name(question):
        return question.item_name + " (" + question.category_name + ").docx"

    @staticmethod
    def generate_single_answer_question_document(question: SingleAnswerQuestion, output_path: str):
        parent_folder = os.path.dirname(os.path.dirname(__file__))
        file_path = os.path.join(parent_folder, "Resources", "UCAT DM Sing Ans Choice Template.docx")
        document: docx.DocumentPart = Document(file_path)
        # Set the item name
        document_table: Table = document.tables[0]
        document_table.cell(0, 2).text = question.item_name
        # Set the category name
        document_table.cell(0, 4).text = question.category_name
        # Set the question stimulus
        WordDocument.add_list_objects_to_cell(document_table.cell(2,1), question.question)
        # Set the question stem
        WordDocument.add_list_objects_to_cell(document_table.cell(4, 1), question.question_stem)
        # Set answers
        WordDocument.add_list_objects_to_cell(document_table.cell(6, 2), question.answers[0].contents)
        WordDocument.add_list_objects_to_cell(document_table.cell(7, 2), question.answers[1].contents)
        WordDocument.add_list_objects_to_cell(document_table.cell(8, 2), question.answers[2].contents)
        WordDocument.add_list_objects_to_cell(document_table.cell(9, 2), question.answers[3].contents)
        # Set correct answer
        WordDocument.set_correct_answer(document_table.cell(10, 2), question.answers)
        # Set the explanation
        WordDocument.add_list_objects_to_cell(document_table.cell(12, 1), question.explanation)
        document.save(output_path)

    @staticmethod
    def set_correct_answer(table: Table, answers: List[AnswerChoice]):
        states = [index for index in range(len(answers)) if
                  answers[index].correct][0]
        if states == 0:
            answer = "A"
        elif states == 1:
            answer = "B"
        elif states == 2:
            answer = "C"
        elif states == 3:
            answer = "D"
        else:
            raise Exception("Correct answer fell outside range A-D")
        table.paragraphs[0].add_run(answer).bold = True

    @staticmethod
    def generate_five_answer_question_document(question: FiveAnswerQuestion, output_path: str):
        parent_folder = os.path.dirname(os.path.dirname(__file__))
        file_path = os.path.join(parent_folder, "Resources", "UCAT DM Five-Part Q Template.docx")
        document: docx.DocumentPart = Document(file_path)
        # Set the item name
        document_table: Table = document.tables[0]
        document_table.cell(0, 2).text = question.item_name
        # Set the category name
        document_table.cell(0, 4).text = question.category_name
        # Set the question stimulus
        WordDocument.add_list_objects_to_cell(document_table.cell(2,1), question.question)
        # Set the question stem
        WordDocument.add_list_objects_to_cell(document_table.cell(4, 1), question.question_stem)
        # Set answers
        WordDocument.set_multi_answer(document_table, 6, question.multi_answers[0])
        WordDocument.set_multi_answer(document_table, 7, question.multi_answers[1])
        WordDocument.set_multi_answer(document_table, 8, question.multi_answers[2])
        WordDocument.set_multi_answer(document_table, 9, question.multi_answers[3])
        WordDocument.set_multi_answer(document_table, 10, question.multi_answers[4])
        # Set the explanation
        WordDocument.add_list_objects_to_cell(document_table.cell(12, 1), question.explanation)
        document.save(output_path)

    @staticmethod
    def set_multi_answer(table: Table, row: int, answer: MultiYesNoAnswer):
        WordDocument.add_list_objects_to_cell(table.cell(row, 2), answer.contents)
        if answer.answer:
            table.cell(row, 5).text = "Yes"
        else:
            table.cell(row, 5).text = "No"

    @staticmethod
    def add_list_objects_to_cell(cell: Table, objects: List[object]):
        paragraph_number = 0
        for i in range(0, len(objects)):
            current_type = type(objects[i])
            if current_type is str:
                cell.paragraphs[paragraph_number].add_run(text=objects[i])
            elif current_type is Graphic:
                cell.add_paragraph()
                paragraph_number += 1
                cell.paragraphs[paragraph_number].add_run().add_picture(image_path_or_stream=objects[i].path,
                                                                               width=Inches(5))
                cell.add_paragraph()
                paragraph_number += 1
            elif current_type is Emphasis:
                if objects[i].bold == True:
                    cell.paragraphs[0].add_run(f" {objects[i].text} ").bold = True
                elif objects[i].italic == True:
                    cell.paragraphs[0].add_run(f" {objects[i].text} ").italic = True
                elif objects[i].underline == True:
                    cell.paragraphs[0].add_run(f" {objects[i].text} ").underline = True
            else:
                raise Exception(f"Unable to determine how to process type {current_type}")
