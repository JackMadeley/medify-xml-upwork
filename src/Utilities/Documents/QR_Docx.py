import os

from docx import Document
from docx.table import Table

from Utilities.Documents.Cell import Cell
from Utilities.Documents.Docx import Docx
from Models.Questions.QRQuantQuestion import QRQuantQuestion
from Models.Questions.QRSimpleQuestion import QRSimpleQuestion


class QR_Docx(Docx):

    def __init__(self):
        super().__init__()
        self.quant_template = os.path.join(self.template_folder, "QR", "UCAT QR Quant Template.docx")
        self.simple_template = os.path.join(self.template_folder, "AR", "UCAT AR Type 2 & 3 Template.docx")

    def generate_quant_template(self, question: QRQuantQuestion):
        if type(question) != QRQuantQuestion:
            raise Exception(f"A QRQuantQuestion was expected but a {type(question)} was provided")
        document = Document(self.quant_template)
        document_table: Table = document.tables[0]
        # Set the item name
        document_table.cell(0, 2).text = question.item_name
        # Set the category name
        document_table.cell(0, 5).text = question.category
        # Add stimulus
        stimulus = Cell(document_table.cell(2, 1))
        stimulus.add_contents_to_cells(question.stimulus)
        # Add question set members
        full_table = Cell(document_table)
        full_table.insert_question_set_members_5_answer(3, question.question_set_members)
        return document

    def generate_simple_template(self, question: QRSimpleQuestion):
        if type(question) != QRSimpleQuestion:
            raise Exception(f"A QRSimpleQuestion was expected but a {type(question)} was provided")
        document = Document(self.simple_template)
        # Set the item name
        document_table: Table = document.tables[0]
        document_table.cell(0, 2).text = question.item_name
        # Set the category name
        document_table.cell(0, 4).text = question.category
        # Set the question stimulus
        stimulus = Cell(document_table.cell(2, 1))
        stimulus.add_contents_to_cells(question.stimulus)
        # Set the question stem
        stem = Cell(document_table.cell(4, 1))
        stem.add_contents_to_cells(question.question_stem)
        # Set answers
        answer_1 = Cell(document_table.cell(6, 2))
        answer_1.add_contents_to_cells(question.answer_choice_set[0].contents)
        answer_2 = Cell(document_table.cell(7, 2))
        answer_2.add_contents_to_cells(question.answer_choice_set[1].contents)
        answer_3 = Cell(document_table.cell(8, 2))
        answer_3.add_contents_to_cells(question.answer_choice_set[2].contents)
        answer_4 = Cell(document_table.cell(9, 2))
        answer_4.add_contents_to_cells(question.answer_choice_set[3].contents)
        # Set correct answer
        correct_answer = Cell(document_table.cell(10, 2))
        correct_answer.set_correct_answer(question.answer_choice_set)
        # Set the explanation
        explanation = Cell(document_table.cell(12, 1))
        explanation.add_contents_to_cells(question.explanation)
        return document
