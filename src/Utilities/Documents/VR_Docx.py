import os

from docx import Document
from docx.table import Table

from Models.Questions.VRQuestion import VRQuestion
from Utilities.Documents.Cell import Cell
from Utilities.Documents.Docx import Docx


class VR_Docx(Docx):

    def __init__(self):
        super().__init__()
        self.question_template = os.path.join(self.template_folder, "VR", "UCAT VR Question Template.docx")
        self.state_template = os.path.join(self.template_folder, "VR", "UCAT VR State Template.docx")

    def generate_question_template(self, question: VRQuestion):
        if type(question) != VRQuestion:
            raise Exception(f"A VRQuestion was expected but a {type(question)} was provided")
        if not all(len(x.answer_choice_set) == 4 for x in question.question_set_members):
            raise Exception("One of more of the answer sets does not have 4 choices")
        document = Document(self.question_template)
        document_table: Table = document.tables[0]
        # Set the item name
        document_table.cell(0, 2).text = question.item_name
        # Set the category name
        document_table.cell(0, 5).text = question.category
        # Add stimulus
        stimulus = Cell(document_table.cell(2, 1))
        stimulus.add_contents_to_cells(question.stimulus)
        # Add question set members
        full_table = Cell(document_table)
        full_table.insert_question_set_members_4_answer(3, question.question_set_members)
        return document

    def generate_state_template(self, question: VRQuestion):
        if type(question) != VRQuestion:
            raise Exception(f"A VRQuestion was expected but a {type(question)} was provided")
        if not all(len(x.answer_choice_set) == 3 for x in question.question_set_members):
            raise Exception("One of more of the answer sets does not have 3 choices")
        document = Document(self.state_template)
        document_table: Table = document.tables[0]
        # Set the item name
        document_table.cell(0, 2).text = question.item_name
        # Set the category name
        document_table.cell(0, 5).text = question.category
        # Add stimulus
        stimulus = Cell(document_table.cell(2, 1))
        stimulus.add_contents_to_cells(question.stimulus)
        # Add question set members
        full_table = Cell(document_table)
        full_table.insert_question_set_members_3_answer(3, question.question_set_members)
        return document

