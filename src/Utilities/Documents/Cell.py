from typing import List

from docx.shared import Inches
from docx.table import Table

from Models.Components.AnswerChoice import AnswerChoice
from Models.Components.Emphasis import Emphasis
from Models.Components.Graphic import Graphic
from Models.Components.Para import Para
from Models.Components.QuestionSetMember import QuestionSetMember
from Models.Components.XMLList import XMLList
from Models.Components.YesNoAnswerChoice import YesNoAnswerChoice


class Cell(object):

    def __init__(self, table: Table):
        self.table = table
        self.paragraph_number = 0

    def add_contents_to_cells(self, contents: List[object]) -> None:
        for item in contents:
            current_type = type(item)
            if current_type is str:
                self.table.paragraphs[self.paragraph_number].add_run(text=item)
            elif current_type is Para:
                self.add_contents_to_cells(item.contents)
                self.table.add_paragraph()
                self.paragraph_number += 1
            elif current_type is Graphic:
                self.table.paragraphs[self.paragraph_number].add_run().add_picture(image_path_or_stream=item.path,
                                                                                  width=Inches(3.5))
                self.table.add_paragraph()
                self.paragraph_number += 1
            elif current_type is Emphasis:
                if item.bold == True:
                    self.table.paragraphs[self.paragraph_number].add_run(item.text).bold = True
                elif item.italic == True:
                    self.table.paragraphs[self.paragraph_number].add_run(item.text).italic = True
                elif item.underline == True:
                    self.table.paragraphs[self.paragraph_number].add_run(item.text).underline = True
            elif current_type is XMLList:
                list: XMLList = item
                if list.prefix == "number":
                    counter = 1
                    for list_item in list.list_items:
                        self.table.paragraphs[self.paragraph_number].add_run(text=f"{counter}. ")
                        self.add_contents_to_cells(list_item)
                else:
                    raise Exception(f"Unhandled list prefix type {list.prefix}")
            else:
                raise Exception(f"Unable to determine how to process type {current_type}")

    def set_multi_answer(self, row: int, answer: YesNoAnswerChoice) -> None:
        answer_box = Cell(self.table.cell(row, 2))
        yes_no_box = Cell(self.table.cell(row, 5))
        answer_box.add_contents_to_cells(answer.contents)
        if answer.correct:
            yes_no_box.add_contents_to_cells(["Yes"])
        else:
            yes_no_box.add_contents_to_cells(["No"])

    def set_correct_answer(self, answers: List[AnswerChoice]) -> None:
        states = [index for index in range(len(answers)) if
                  answers[index].correct][0]
        if states == 0:
            answer = "A"
        elif states == 1:
            answer = "B"
        elif states == 2:
            answer = "C"
        elif states == 3:
            answer = "D"
        else:
            raise Exception("Correct answer fell outside range A-D")
        self.table.paragraphs[self.paragraph_number].add_run(answer).bold = True

    def insert_question_set_members_3_answer(self, start_row: int, question_set_members: List[QuestionSetMember]) -> None:
        for member in question_set_members:
            self.insert_question_set_member_3_answer(start_row, member)
            start_row += 10

    def insert_question_set_member_3_answer(self, start_row: int, question_set_member: QuestionSetMember) -> None:
        question_item_name = Cell(self.table.cell(start_row, 3))
        question_item_name.add_contents_to_cells([question_set_member.item_name])
        question_category_name = Cell(self.table.cell(start_row, 6))
        question_category_name.add_contents_to_cells([question_set_member.category])
        question_stem = Cell(self.table.cell(start_row+2, 1))
        question_stem.add_contents_to_cells(question_set_member.question_stem)
        answer_a = Cell(self.table.cell(start_row+4, 2))
        answer_a.add_contents_to_cells(question_set_member.answer_choice_set[0].contents)
        answer_b = Cell(self.table.cell(start_row+5, 2))
        answer_b.add_contents_to_cells(question_set_member.answer_choice_set[1].contents)
        answer_c = Cell(self.table.cell(start_row+6, 2))
        answer_c.add_contents_to_cells(question_set_member.answer_choice_set[2].contents)
        correct_answer = Cell(self.table.cell(start_row+7, 2))
        correct_answer.set_correct_answer(question_set_member.answer_choice_set)
        explanation = Cell(self.table.cell(start_row+9, 1))
        explanation.add_contents_to_cells(question_set_member.explanation)

    def insert_question_set_members_4_answer(self, start_row: int, question_set_members: List[QuestionSetMember]) -> None:
        for member in question_set_members:
            self.insert_question_set_member_4_answer(start_row, member)
            start_row += 11

    def insert_question_set_member_4_answer(self, start_row: int, question_set_member: QuestionSetMember) -> None:
        question_item_name = Cell(self.table.cell(start_row, 3))
        question_item_name.add_contents_to_cells([question_set_member.item_name])
        question_category_name = Cell(self.table.cell(start_row, 6))
        question_category_name.add_contents_to_cells([question_set_member.category])
        question_stem = Cell(self.table.cell(start_row+2, 1))
        question_stem.add_contents_to_cells(question_set_member.question_stem)
        answer_a = Cell(self.table.cell(start_row+4, 2))
        answer_a.add_contents_to_cells(question_set_member.answer_choice_set[0].contents)
        answer_b = Cell(self.table.cell(start_row+5, 2))
        answer_b.add_contents_to_cells(question_set_member.answer_choice_set[1].contents)
        answer_c = Cell(self.table.cell(start_row+6, 2))
        answer_c.add_contents_to_cells(question_set_member.answer_choice_set[2].contents)
        answer_d = Cell(self.table.cell(start_row+7, 2))
        answer_d.add_contents_to_cells(question_set_member.answer_choice_set[3].contents)
        correct_answer = Cell(self.table.cell(start_row+8, 2))
        correct_answer.set_correct_answer(question_set_member.answer_choice_set)
        explanation = Cell(self.table.cell(start_row+10, 1))
        explanation.add_contents_to_cells(question_set_member.explanation)

    def insert_question_set_members_5_answer(self, start_row: int, question_set_members: List[QuestionSetMember]) -> None:
        for member in question_set_members:
            self.insert_question_set_member_5_answer(start_row, member)
            start_row += 12

    def insert_question_set_member_5_answer(self, start_row: int, question_set_member: QuestionSetMember) -> None:
        question_item_name = Cell(self.table.cell(start_row, 3))
        question_item_name.add_contents_to_cells([question_set_member.item_name])
        question_category_name = Cell(self.table.cell(start_row, 6))
        question_category_name.add_contents_to_cells([question_set_member.category])
        question_stem = Cell(self.table.cell(start_row+2, 1))
        question_stem.add_contents_to_cells(question_set_member.question_stem)
        answer_a = Cell(self.table.cell(start_row+4, 2))
        answer_a.add_contents_to_cells(question_set_member.answer_choice_set[0].contents)
        answer_b = Cell(self.table.cell(start_row+5, 2))
        answer_b.add_contents_to_cells(question_set_member.answer_choice_set[1].contents)
        answer_c = Cell(self.table.cell(start_row+6, 2))
        answer_c.add_contents_to_cells(question_set_member.answer_choice_set[2].contents)
        answer_d = Cell(self.table.cell(start_row+7, 2))
        answer_d.add_contents_to_cells(question_set_member.answer_choice_set[3].contents)
        answer_e = Cell(self.table.cell(start_row+8, 2))
        answer_e.add_contents_to_cells(question_set_member.answer_choice_set[4].contents)
        correct_answer = Cell(self.table.cell(start_row+9, 2))
        correct_answer.set_correct_answer(question_set_member.answer_choice_set)
        explanation = Cell(self.table.cell(start_row+11, 1))
        explanation.add_contents_to_cells(question_set_member.explanation)