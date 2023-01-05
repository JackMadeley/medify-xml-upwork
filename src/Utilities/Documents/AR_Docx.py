import os

from docx import Document
from docx.table import Table

from Utilities.Documents.Cell import Cell
from Utilities.Documents.Docx import Docx
from Models.Questions.ARType1Question import ARType1Question
from Models.Questions.ARType2Question import ARType2Question
from Models.Questions.ARType3Question import ARType3Question
from Models.Questions.ARType4Question import ARType4Question


class AR_Docx(Docx):

    def __init__(self):
        super().__init__()
        self.type_1_template = os.path.join(self.template_folder, "AR", "UCAT AR Type 1 Template.docx")
        self.type_2_3_template = os.path.join(self.template_folder, "AR", "UCAT AR Type 2 & 3 Template.docx")
        self.type_4_template = os.path.join(self.template_folder, "AR", "UCAT AR Type 4 Template.docx")

    def generate_type_1_template(self, question: ARType1Question):
        if type(question) != ARType1Question:
            raise Exception(f"A ARType1Question was expected but a {type(question)} was provided")
        document = Document(self.type_1_template)
        document_table: Table = document.tables[0]
        # Set the item name
        document_table.cell(0, 2).text = question.item_name
        # Set the category name
        document_table.cell(0, 5).text = question.category
        # Add stimulus
        stimulus = Cell(document_table.cell(2, 1))
        stimulus.add_contents_to_cells(question.stimulus)
        # Add question set members
        full_table = Cell(document_table)
        full_table.insert_question_set_members_3_answer(3, question.question_set_members)
        return document

    def generate_type_2_template(self, question: ARType2Question):
        if type(question) != ARType2Question:
            raise Exception(f"A ARType2Question was expected but a {type(question)} was provided")
        document = Document(self.type_2_3_template)
        document_table: Table = document.tables[0]
        document_table.cell(0, 2).text = question.item_name
        # Set the category name
        document_table.cell(0, 4).text = question.category
        # Set the question stimulus
        stimulus = Cell(document_table.cell(2, 1))
        stimulus.add_contents_to_cells(question.stimulus)
        # Set the question stem
        stem = Cell(document_table.cell(4, 1))
        stem.add_contents_to_cells(question.question_stem)
        # Set answers
        answer_1 = Cell(document_table.cell(6, 2))
        answer_1.add_contents_to_cells(question.answer_choice_set[0].contents)
        answer_2 = Cell(document_table.cell(7, 2))
        answer_2.add_contents_to_cells(question.answer_choice_set[1].contents)
        answer_3 = Cell(document_table.cell(8, 2))
        answer_3.add_contents_to_cells(question.answer_choice_set[2].contents)
        answer_4 = Cell(document_table.cell(9, 2))
        answer_4.add_contents_to_cells(question.answer_choice_set[3].contents)
        # Set correct answer
        correct_answer = Cell(document_table.cell(10, 2))
        correct_answer.set_correct_answer(question.answer_choice_set)
        # Set the explanation
        explanation = Cell(document_table.cell(12, 1))
        explanation.add_contents_to_cells(question.explanation)
        return document

    def generate_type_3_template(self, question: ARType3Question):
        if type(question) != ARType3Question:
            raise Exception(f"A ARType3Question was expected but a {type(question)} was provided")
        document = Document(self.type_2_3_template)
        document_table: Table = document.tables[0]
        document_table.cell(0, 2).text = question.item_name
        # Set the category name
        document_table.cell(0, 4).text = question.category
        # Set the question stimulus
        stimulus = Cell(document_table.cell(2, 1))
        stimulus.add_contents_to_cells(question.question)
        # Set the question stem
        stem = Cell(document_table.cell(4, 1))
        stem.add_contents_to_cells(question.question_stem)
        # Set answers
        answer_1 = Cell(document_table.cell(6, 2))
        answer_1.add_contents_to_cells(question.answer_choice_set[0].contents)
        answer_2 = Cell(document_table.cell(7, 2))
        answer_2.add_contents_to_cells(question.answer_choice_set[1].contents)
        answer_3 = Cell(document_table.cell(8, 2))
        answer_3.add_contents_to_cells(question.answer_choice_set[2].contents)
        answer_4 = Cell(document_table.cell(9, 2))
        answer_4.add_contents_to_cells(question.answer_choice_set[3].contents)
        # Set correct answer
        correct_answer = Cell(document_table.cell(10, 2))
        correct_answer.set_correct_answer(question.answer_choice_set)
        # Set the explanation
        explanation = Cell(document_table.cell(12, 1))
        explanation.add_contents_to_cells(question.explanation)
        return document

    def generate_type_4_template(self, question: ARType4Question):
        if type(question) != ARType4Question:
            raise Exception(f"A ARType4Question was expected but a {type(question)} was provided")
        document = Document(self.type_4_template)
        document_table: Table = document.tables[0]
        # Set the item name
        document_table.cell(0, 2).text = question.item_name
        # Set the category name
        document_table.cell(0, 5).text = question.category
        # Add stimulus
        stimulus = Cell(document_table.cell(2, 1))
        stimulus.add_contents_to_cells(question.stimulus)
        # Add question set members
        full_table = Cell(document_table)
        full_table.insert_question_set_members_4_answer(3, question.question_set_members)
        return document
