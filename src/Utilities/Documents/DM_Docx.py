import os

from docx import Document
from docx.table import Table

from Utilities.Documents.Cell import Cell
from Utilities.Documents.Docx import Docx
from Models.Questions.DMFiveAnswerQuestion import DMFiveAnswerQuestion
from Models.Questions.DMSingleAnswerQuestion import DMSingleAnswerQuestion


class DM_Docx(Docx):

    def __init__(self):
        super().__init__()
        self.single_part_template = os.path.join(self.template_folder, "DM", "UCAT DM Sing Ans Choice Template.docx")
        self.five_part_template = os.path.join(self.template_folder, "DM", "UCAT DM Five-Part Q Template.docx")

    def generate_five_part_question(self, question: DMFiveAnswerQuestion):
        if type(question) != DMFiveAnswerQuestion:
            raise Exception(f"A DMFiveAnswerQuestion was expected but a {type(question)} was provided")
        document = Document(self.five_part_template)
        # Set the item name
        document_table: Table = document.tables[0]
        document_table.cell(0, 2).text = question.item_name
        # Set the category name
        document_table.cell(0, 4).text = question.category
        # Set the question stimulus
        stimulus = Cell(document_table.cell(2, 1))
        stimulus.add_contents_to_cells(question.question)
        # Set the question stem
        stem = Cell(document_table.cell(4, 1))
        stem.add_contents_to_cells(question.question_stem)
        # Set answers
        answer_table = Cell(document_table)
        answer_table.set_multi_answer(6, question.multi_answers_set[0])
        answer_table.set_multi_answer(7, question.multi_answers_set[1])
        answer_table.set_multi_answer(8, question.multi_answers_set[2])
        answer_table.set_multi_answer(9, question.multi_answers_set[3])
        answer_table.set_multi_answer(10, question.multi_answers_set[4])
        # Set the explanation
        explanation = Cell(document_table.cell(12, 1))
        explanation.add_contents_to_cells(question.explanation)
        return document

    def generate_single_part_question(self, question: DMSingleAnswerQuestion):
        if type(question) != DMSingleAnswerQuestion:
            raise Exception(f"A DMSingleAnswerQuestion was expected but a {type(question)} was provided")
        document = Document(self.single_part_template)
        # Set the item name
        document_table: Table = document.tables[0]
        document_table.cell(0, 2).text = question.item_name
        # Set the category name
        document_table.cell(0, 4).text = question.category
        # Set the question stimulus
        stimulus = Cell(document_table.cell(2, 1))
        stimulus.add_contents_to_cells(question.question)
        # Set the question stem
        stem = Cell(document_table.cell(4, 1))
        stem.add_contents_to_cells(question.question_stem)
        # Set answers
        answer_1 = Cell(document_table.cell(6, 2))
        answer_1.add_contents_to_cells(question.answer_choice_set[0].contents)
        answer_2 = Cell(document_table.cell(7, 2))
        answer_2.add_contents_to_cells(question.answer_choice_set[1].contents)
        answer_3 = Cell(document_table.cell(8, 2))
        answer_3.add_contents_to_cells(question.answer_choice_set[2].contents)
        answer_4 = Cell(document_table.cell(9, 2))
        answer_4.add_contents_to_cells(question.answer_choice_set[3].contents)
        # Set correct answer
        correct_answer = Cell(document_table.cell(10, 2))
        correct_answer.set_correct_answer(question.answer_choice_set)
        # Set the explanation
        explanation = Cell(document_table.cell(12, 1))
        explanation.add_contents_to_cells(question.explanation)
        return document
